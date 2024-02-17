import neat
from qiskit import Aer, QuantumCircuit, execute
import cupy as cp

# Configuration file path
config_file = "CUDAANN-neat-config.txt"  # Update this path

# Load NEAT configuration
config = neat.Config(neat.DefaultGenome, neat.DefaultReproduction,
                     neat.DefaultSpeciesSet, neat.DefaultStagnation,
                     config_file)

# Create a population from the config
population = neat.Population(config)

# Add reporters to show progress in the terminal
population.add_reporter(neat.StdOutReporter(True))
population.add_reporter(neat.StatisticsReporter())

# Quantum operation function
def perform_quantum_operation(input_data):
    circuit = QuantumCircuit(len(input_data), len(input_data))
    # Example: Apply Hadamard gate to all qubits to create superposition
    for i in range(len(input_data)):
        circuit.h(i)
    circuit.measure(range(len(input_data)), range(len(input_data)))
    
    backend = Aer.get_backend('qasm_simulator')
    job = execute(circuit, backend, shots=1)
    result = job.result().get_counts()
    
    # Simplified result processing: Convert to binary list
    measured_result = list(result.keys())[0]
    return [int(bit) for bit in measured_result]

# Fitness evaluation function
def eval_genomes(genomes, config):
    for genome_id, genome in genomes:
        net = neat.nn.FeedForwardNetwork.create(genome, config)
        
        # Example: Prepare data for quantum operation
        quantum_data = [0, 1]  # Placeholder data
        
        # Perform quantum operation and use the result as input
        quantum_result = perform_quantum_operation(quantum_data)
        network_input = cp.asarray(quantum_result).get()  # Convert to numpy array
        
        # Evaluate the network output to calculate fitness
        output = net.activate(network_input)
        fitness = output[0]  # Example: Use the first output to determine fitness
        
        genome.fitness = fitness

# Run the evolution
winner = population.run(eval_genomes, 50)  # Adjust the number of generations if needed

print(f"Winner: {winner}")


# This function call is just an example; you'll need to replace 'network' with the actual network object from NEAT
# visualize_network(best_network)

from sklearn.feature_extraction.text import TfidfVectorizer

# Example text data
texts = ["This is a sample sentence", "Another example text", "Neural networks can learn from text"]

# Convert text data to TF-IDF vectors
vectorizer = TfidfVectorizer(max_features=500)  # Limiting to 500 features for simplicity
X = vectorizer.fit_transform(texts).toarray()

# Assuming binary classification for simplicity
y = [0, 1, 0]  # Example target labels for each text

import cupy as cp

# Placeholder for network initialization
# This needs to be adapted based on the winning genome's structure
num_inputs = 500  # Should match the number of features in your text representation
num_outputs = 1  # For binary classification
# Initialize weights and biases for each layer based on the winning genome
weights = cp.random.randn(num_inputs, num_outputs)
biases = cp.random.randn(num_outputs)

# Simple forward pass
import cupy as cp

# Example: Simple forward pass through one layer of the network
def forward_pass(input_vector, weights):
    # Assuming input_vector and weights are CuPy arrays
    activation = cp.dot(weights, input_vector)
    output = cp.tanh(activation)  # Example activation function
    return output

# Note: You'll need to translate your NEAT network architecture into CuPy operations,
# including initializing weights and handling activations.
from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np

# Example preprocessing steps
def preprocess_texts(texts):
    # Assume `texts` is a list of strings extracted from DOCX documents
    vectorizer = TfidfVectorizer(max_features=500, stop_words='english')
    X = vectorizer.fit_transform(texts).toarray()
    return X

# Example mapping based on hypothetical network output
dds_mapping = {
    0: '000 – Computer science, information & general works',
    1: '100 – Philosophy & psychology',
    # Add mappings for other categories
}

def classify_and_map_to_dds(network_output):
    # Assuming `network_output` is the index of the highest probability output
    dds_category = dds_mapping.get(network_output, 'Undefined')
    return dds_category

import zipfile
import os

def extract_images(docx_path, output_folder):
    with zipfile.ZipFile(docx_path, 'r') as docx:
        for item in docx.infolist():
            if item.filename.startswith('word/media/'):
                image_data = docx.read(item.filename)
                output_path = os.path.join(output_folder, os.path.basename(item.filename))
                with open(output_path, 'wb') as image_file:
                    image_file.write(image_data)

from keras.preprocessing.image import ImageDataGenerator

# Image dimensions and batch size
img_height, img_width = 224, 224
batch_size = 32

datagen = ImageDataGenerator(rescale=1./255, validation_split=0.2)

train_generator = datagen.flow_from_directory(
    'C:/Users/HeadAdminKiriguya/Documents/DOCX/',
    target_size=(img_height, img_width),
    batch_size=batch_size,
    class_mode='binary',  # for binary classification
    subset='training')

validation_generator = datagen.flow_from_directory(
    'C:/Users/HeadAdminKiriguya/Documents/DOCX/',
    target_size=(img_height, img_width),
    batch_size=batch_size,
    class_mode='binary',  # for binary classification
    subset='validation')

from keras.applications import MobileNetV2
from keras.models import Model
from keras.layers import Dense, GlobalAveragePooling2D

# Load the base model
base_model = MobileNetV2(weights='imagenet', include_top=False, input_shape=(img_height, img_width, 3))

# Freeze the base model
base_model.trainable = False

# Add custom layers
x = base_model.output
x = GlobalAveragePooling2D()(x)
x = Dense(1024, activation='relu')(x)
predictions = Dense(1, activation='sigmoid')(x)

# Create the final model
model = Model(inputs=base_model.input, outputs=predictions)

# Compile the model
model.compile(optimizer='adam', loss='binary_crossentropy', metrics=['accuracy'])

print(f"Found {train_generator.samples} training samples.")
print(f"Found {validation_generator.samples} validation samples.")


validation_steps = max(1, validation_generator.samples // batch_size)



model.save('equation_detector.h5')




from docx import Document
import re

def clean_text(text):
    """
    Cleans the input text by lowering case, removing punctuation,
    and stripping whitespace.

    Parameters:
    - text: str, original text to clean.

    Returns:
    - str, cleaned text.
    """
    # Lower case
    text = text.lower()
    
    # Remove punctuation
    text = re.sub(r'[^\w\s]', '', text)
    
    # Remove numbers (optional, depending on your needs)
    text = re.sub(r'\d+', '', text)
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def extract_text_from_docx(file_path):
    """
    Extracts text from a DOCX file.

    Parameters:
    - file_path: str, the path to the DOCX file.

    Returns:
    - str, the extracted text from the DOCX file.
    """
    # Ensure that file_path is a string representing the path to a single DOCX file
    doc = Document(file_path)  # Correct usage: passing a single file path
    full_text = [para.text for para in doc.paragraphs]
    return '\n'.join(full_text)

    # Example usage:
doc_paths = doc_paths = [
    "C:\\Users\\HeadAdminKiriguya\\Documents\\DOCX\\Holographicxrinterface.docx",
    "C:\\Users\\HeadAdminKiriguya\\Documents\\DOCX\\Holographicpolarizedfield.docx",]
    # Make sure each path is correct

texts = [extract_text_from_docx(doc_path) for doc_path in doc_paths]  
# Assuming the use of previously defined functions for extraction and preprocessing
texts = [extract_text_from_docx(doc_path) for doc_path in doc_paths]
cleaned_texts = [clean_text(text) for text in texts]

vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
X = vectorizer.fit_transform(cleaned_texts).toarray()

from sklearn.cluster import KMeans
import matplotlib.pyplot as plt


# Assuming `texts` contains the extracted texts from your DOCX documents
cleaned_texts = [clean_text(text) for text in texts]

# Using the Elbow Method to find the optimal number of clusters
wcss = []
for i in range(1, 11):
    # Adjusting n_clusters to match the number of samples
    kmeans = KMeans(n_clusters=2, init='k-means++', random_state=42)
    kmeans.fit(X)

    wcss.append(kmeans.inertia_)

plt.plot(range(1, 11), wcss)
plt.title('Elbow Method')
plt.xlabel('Number of clusters')
plt.ylabel('WCSS')  # Within cluster sum of squares
plt.show()

# Choose the number of clusters (k) based on the plot
# Assuming X is your data
n_samples = X.shape[0]
n_clusters = min(n_samples, 5)  # Ensure n_clusters does not exceed the number of samples

# Example of defining k directly
k = min(X.shape[0], 5)  # Assuming X is your dataset

# Alternatively, after fitting the model, you can retrieve the number of clusters like this:
kmeans = KMeans(n_clusters=k, init='k-means++', random_state=42)
y_kmeans = kmeans.fit_predict(X)
y_kmeans = kmeans.fit_predict(X)
kmeans.fit(X)
# No need to redefine 'k' here since it's already set

# Now you can use 'k' in your loop without issue
for i in range(k):
    # your loop logic

    print(f"Cluster {i}:")
    cluster_docs = [texts[j] for j in range(len(texts)) if y_kmeans[j] == i]
    # Optionally print some representative documents or titles
    print(cluster_docs[:5])  # Print first 5 documents' text for review


from docx import Document

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return " ".join([para.text for para in doc.paragraphs])

import re
from sklearn.feature_extraction.text import TfidfVectorizer

def clean_text(text):
    text = text.lower()
    text = re.sub(r'\[.*?\]', '', text)
    text = re.sub(r'[%s]' % re.escape('!"#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~'), '', text)
    text = re.sub(r'\w*\d\w*', '', text)
    return text

# Example usage with a list of extracted texts
texts = [extract_text_from_docx(doc_path) for doc_path in doc_paths]
cleaned_texts = [clean_text(text) for text in texts]

vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
X = vectorizer.fit_transform(cleaned_texts)

from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
import os

# Example: Assuming doc_paths contains file paths
# and you want to use the parent directory name as the label
y = [os.path.basename(os.path.dirname(doc_path)) for doc_path in doc_paths]
# Assuming kmeans.fit_predict(X) was used and X corresponds to features extracted from documents in doc_paths
cluster_labels = kmeans.fit_predict(X)

# Associate each document path with its cluster label
y = [cluster_labels[i] for i in range(len(doc_paths))]

# Defining label directly
label = "Dataset"

# Now you can use label in your list comprehension
y = [label for doc_path in doc_paths]

# Defining labels based on some function of doc_path
def determine_label(doc_path):
    # Implement logic to determine the label
    return "DeterminedLabel"

# Use determine_label to assign labels
y = [determine_label(doc_path) for doc_path in doc_paths]


X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

model = RandomForestClassifier(n_estimators=100)
model.fit(X_train, y_train)

from sklearn.metrics import classification_report

y_pred = model.predict(X_test)
print(classification_report(y_test, y_pred))

from joblib import dump

dump(model, 'text_classification_model.joblib')
